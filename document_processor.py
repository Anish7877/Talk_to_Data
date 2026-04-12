"""
Document Processor
Handles uploaded files and routes them to the correct pipeline:
  - Structured data (CSV, Excel, JSON) → PostgreSQL + TAG schema
  - Unstructured docs (PDF, TXT, DOCX, MD) → ChromaDB RAG collection

Usage:
    from document_processor import DocumentProcessor
    processor = DocumentProcessor(tag=tag_instance, executor=executor_instance)
    result = processor.process("path/to/file.csv")
    result = processor.process("path/to/report.pdf")
"""

import os
import json
import uuid
import logging
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# File type classification
# ---------------------------------------------------------------------------

STRUCTURED_EXTENSIONS   = {".csv", ".xlsx", ".xls", ".json"}
UNSTRUCTURED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".markdown"}


def classify_file(file_path: str) -> str:
    """
    Classify a file as 'structured' or 'unstructured'.

    Returns:
        'structured'   → CSV, Excel, JSON  (goes to Postgres + TAG)
        'unstructured' → PDF, TXT, DOCX, MD (goes to ChromaDB RAG)
        'unsupported'  → anything else
    """
    ext = Path(file_path).suffix.lower()
    if ext in STRUCTURED_EXTENSIONS:
        return "structured"
    if ext in UNSTRUCTURED_EXTENSIONS:
        return "unstructured"
    return "unsupported"


# ---------------------------------------------------------------------------
# Structured data handlers
# ---------------------------------------------------------------------------

class StructuredFileLoader:
    """Loads structured files (CSV, Excel, JSON) into a pandas DataFrame."""

    def load(self, file_path: str):
        """
        Load any supported structured file into a DataFrame.
        Returns (DataFrame, table_name_suggestion).
        """
        import pandas as pd

        ext  = Path(file_path).suffix.lower()
        stem = Path(file_path).stem.lower().replace(" ", "_").replace("-", "_")

        if ext == ".csv":
            df = pd.read_csv(file_path)
        elif ext in {".xlsx", ".xls"}:
            df = pd.read_excel(file_path)
        elif ext == ".json":
            df = pd.read_json(file_path)
        else:
            raise ValueError(f"Unsupported structured format: {ext}")

        # Sanitize column names
        df.columns = [
            c.lower().strip().replace(" ", "_").replace("-", "_")
            for c in df.columns
        ]

        logger.info(f"Loaded {len(df)} rows, {len(df.columns)} columns from {file_path}")
        return df, stem

    def infer_postgres_type(self, dtype) -> str:
        """Map pandas dtype to PostgreSQL type."""
        import pandas as pd
        dtype_str = str(dtype)
        if "int" in dtype_str:
            return "BIGINT"
        if "float" in dtype_str:
            return "DOUBLE PRECISION"
        if "bool" in dtype_str:
            return "BOOLEAN"
        if "datetime" in dtype_str:
            return "TIMESTAMP"
        if "date" in dtype_str:
            return "DATE"
        return "TEXT"

    def create_table_and_insert(
        self,
        df,
        table_name: str,
        engine,           # SQLAlchemy engine (admin connection, not read-only)
        if_exists: str = "replace"
    ) -> bool:
        """
        Write DataFrame to PostgreSQL and grant read access to ai_readonly.
        Uses pandas to_sql for simplicity.
        """
        from sqlalchemy import text

        try:
            # Write table (pandas handles CREATE TABLE automatically)
            df.to_sql(
                table_name,
                engine,
                if_exists=if_exists,
                index=False,
                method="multi",
                chunksize=500
            )

            # Grant read access to the read-only role
            with engine.connect() as conn:
                conn.execute(text(
                    f"GRANT SELECT ON TABLE {table_name} TO ai_readonly;"
                ))
                conn.commit()

            logger.info(f"Table '{table_name}' created/updated with {len(df)} rows")
            return True

        except Exception as e:
            logger.error(f"Failed to write table '{table_name}': {e}")
            return False

    def build_table_description(self, df, table_name: str, file_path: str):
        """
        Build a TableDescription object for the TAG layer so the
        SQL agent knows this table exists.
        """
        from layers.layer3_tag import TableDescription

        columns = []
        sample_values = {}

        for col in df.columns:
            pg_type = self.infer_postgres_type(df[col].dtype)
            columns.append({
                "name":        col,
                "type":        pg_type,
                "description": f"Column from uploaded file {Path(file_path).name}"
            })
            # Pick a non-null sample value
            non_null = df[col].dropna()
            if len(non_null) > 0:
                sample_values[col] = str(non_null.iloc[0])

        return TableDescription(
            table_name=table_name,
            description=(
                f"Uploaded structured dataset from '{Path(file_path).name}'. "
                f"Contains {len(df)} rows and {len(df.columns)} columns."
            ),
            columns=columns,
            relationships=[],
            sample_values=sample_values
        )


# ---------------------------------------------------------------------------
# Unstructured document handlers
# ---------------------------------------------------------------------------

class UnstructuredFileLoader:
    """Extracts text from PDF, TXT, DOCX, MD files."""

    def load(self, file_path: str) -> str:
        """Extract raw text from the file. Returns text string."""
        ext = Path(file_path).suffix.lower()

        if ext == ".txt" or ext in {".md", ".markdown"}:
            return self._load_text(file_path)
        elif ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported unstructured format: {ext}")

    def _load_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf (install: pip install pypdf)."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF support. Install it with:\n"
                "  pip install pypdf"
            )

    def _load_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx (install: pip install python-docx)."""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. Install it with:\n"
                "  pip install python-docx"
            )

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[str]:
        """
        Split text into overlapping chunks for embedding.
        Simple word-boundary chunking.
        """
        words  = text.split()
        chunks = []
        start  = 0

        while start < len(words):
            end   = min(start + chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk)
            if end == len(words):
                break
            start = end - overlap  # overlap for context continuity

        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks


# ---------------------------------------------------------------------------
# Main DocumentProcessor
# ---------------------------------------------------------------------------

class DocumentProcessor:
    """
    Main entry point for file processing.

    Requires:
        tag      - TAGRetrieval instance (Layer 3)
        executor - SecureExecutionSandbox instance (Layer 5) for test_connection()

    For structured files, also needs:
        admin_engine - SQLAlchemy engine with write permissions (NOT the read-only one)
    """

    def __init__(
        self,
        tag,
        executor=None,
        admin_db_url: str = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        """
        Args:
            tag:           TAGRetrieval instance
            executor:      SecureExecutionSandbox (used only for health check)
            admin_db_url:  SQLAlchemy URL with write access for structured uploads
                           e.g. "postgresql://postgres:password@localhost/postgres"
            chunk_size:    Words per chunk for RAG
            chunk_overlap: Overlapping words between chunks
        """
        self.tag          = tag
        self.executor     = executor
        self.chunk_size   = chunk_size
        self.chunk_overlap = chunk_overlap

        self._structured_loader   = StructuredFileLoader()
        self._unstructured_loader = UnstructuredFileLoader()

        # Admin engine (write access) for structured uploads
        self._admin_engine = None
        if admin_db_url:
            from sqlalchemy import create_engine
            self._admin_engine = create_engine(admin_db_url, pool_pre_ping=True)
            logger.info("Admin DB engine initialized for structured uploads")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Main method. Auto-detects file type and processes accordingly.

        Returns a result dict:
        {
            "success":   bool,
            "file_type": "structured" | "unstructured" | "unsupported",
            "file_name": str,
            "message":   str,
            # structured only:
            "table_name": str,
            "row_count":  int,
            "columns":    List[str],
            # unstructured only:
            "chunk_count": int,
            "doc_ids":     List[str],
        }
        """
        file_path = str(file_path)
        file_name = Path(file_path).name
        file_type = classify_file(file_path)

        logger.info(f"Processing '{file_name}' as {file_type}")

        if file_type == "structured":
            return self._process_structured(file_path, file_name)
        elif file_type == "unstructured":
            return self._process_unstructured(file_path, file_name)
        else:
            return {
                "success":   False,
                "file_type": "unsupported",
                "file_name": file_name,
                "message":   (
                    f"Unsupported file type: {Path(file_path).suffix}. "
                    f"Supported: {STRUCTURED_EXTENSIONS | UNSTRUCTURED_EXTENSIONS}"
                )
            }

    def process_many(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple files. Returns list of result dicts."""
        return [self.process(p) for p in file_paths]

    def list_loaded_schemas(self) -> List[str]:
        """List all table names currently in the TAG schema collection."""
        try:
            count = self.tag.schema_collection.count()
            if count == 0:
                return []
            results = self.tag.schema_collection.get()
            return results.get("ids", [])
        except Exception:
            return []

    def list_loaded_documents(self) -> List[Dict[str, str]]:
        """List all document chunks in the TAG docs collection."""
        try:
            count = self.tag.docs_collection.count()
            if count == 0:
                return []
            results = self.tag.docs_collection.get()
            ids       = results.get("ids", [])
            metadatas = results.get("metadatas", [{}] * len(ids))
            return [
                {"id": id_, "file_name": m.get("file_name", "unknown")}
                for id_, m in zip(ids, metadatas)
            ]
        except Exception:
            return []

    # ------------------------------------------------------------------
    # Structured processing
    # ------------------------------------------------------------------

    def _process_structured(self, file_path: str, file_name: str) -> Dict[str, Any]:
        try:
            # 1. Load into DataFrame
            df, table_name = self._structured_loader.load(file_path)

            # 2. Write to Postgres (if admin engine available)
            db_written = False
            if self._admin_engine:
                db_written = self._structured_loader.create_table_and_insert(
                    df, table_name, self._admin_engine
                )
                if not db_written:
                    logger.warning(
                        f"Could not write '{table_name}' to DB. "
                        "Schema will still be added to TAG so the LLM knows about it."
                    )
            else:
                logger.warning(
                    "No admin_db_url provided — skipping DB write. "
                    "Schema will be added to TAG only."
                )

            # 3. Add schema to TAG so SQL agent can query it
            table_desc = self._structured_loader.build_table_description(
                df, table_name, file_path
            )
            self.tag.add_schema(table_desc)

            return {
                "success":    True,
                "file_type":  "structured",
                "file_name":  file_name,
                "table_name": table_name,
                "row_count":  len(df),
                "columns":    list(df.columns),
                "db_written": db_written,
                "message":    (
                    f"Loaded {len(df)} rows into table '{table_name}' "
                    f"and added schema to TAG."
                    + ("" if db_written else " (DB write skipped — no admin connection)")
                )
            }

        except Exception as e:
            logger.error(f"Structured processing failed for '{file_name}': {e}")
            return {
                "success":   False,
                "file_type": "structured",
                "file_name": file_name,
                "message":   f"Failed: {str(e)}"
            }

    # ------------------------------------------------------------------
    # Unstructured processing
    # ------------------------------------------------------------------

    def _process_unstructured(self, file_path: str, file_name: str) -> Dict[str, Any]:
        try:
            # 1. Extract text
            text = self._unstructured_loader.load(file_path)

            if not text.strip():
                return {
                    "success":   False,
                    "file_type": "unstructured",
                    "file_name": file_name,
                    "message":   "File appears to be empty or could not extract text."
                }

            # 2. Chunk text
            chunks = self._unstructured_loader.chunk_text(
                text,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )

            # 3. Store each chunk in ChromaDB docs collection
            # Use file hash + chunk index as stable ID
            file_hash = hashlib.md5(file_name.encode()).hexdigest()[:8]
            doc_ids   = []

            for i, chunk in enumerate(chunks):
                doc_id = f"{file_hash}_chunk_{i}"
                self.tag.add_document(
                    doc_id=doc_id,
                    content=chunk,
                    metadata={
                        "file_name":  file_name,
                        "file_path":  file_path,
                        "chunk_index": str(i),
                        "total_chunks": str(len(chunks)),
                        "file_type":  Path(file_path).suffix.lower()
                    }
                )
                doc_ids.append(doc_id)

            return {
                "success":     True,
                "file_type":   "unstructured",
                "file_name":   file_name,
                "chunk_count": len(chunks),
                "doc_ids":     doc_ids,
                "char_count":  len(text),
                "message":     (
                    f"Extracted {len(text):,} characters from '{file_name}', "
                    f"split into {len(chunks)} chunks, stored in RAG collection."
                )
            }

        except Exception as e:
            logger.error(f"Unstructured processing failed for '{file_name}': {e}")
            return {
                "success":   False,
                "file_type": "unstructured",
                "file_name": file_name,
                "message":   f"Failed: {str(e)}"
            }


# ---------------------------------------------------------------------------
# Factory function
# ---------------------------------------------------------------------------

def create_document_processor(
    tag,
    executor=None,
    config: Dict[str, Any] = None
) -> DocumentProcessor:
    """
    Create a DocumentProcessor from config.

    config keys used:
        db_host, db_port, db_name, db_user (admin), db_password (admin)
        chunk_size, chunk_overlap
    """
    config = config or {}

    # Build admin DB URL from config (uses postgres superuser, not ai_readonly)
    admin_db_url = None
    db_host     = config.get("db_host", "localhost")
    db_port     = config.get("db_port", 5432)
    db_name     = config.get("db_name", "postgres")
    admin_user  = config.get("admin_db_user", config.get("db_user", "postgres"))
    admin_pass  = config.get("admin_db_password", config.get("db_password", ""))

    if admin_user and admin_pass:
        admin_db_url = f"postgresql://{admin_user}:{admin_pass}@{db_host}:{db_port}/{db_name}"

    return DocumentProcessor(
        tag=tag,
        executor=executor,
        admin_db_url=admin_db_url,
        chunk_size=config.get("chunk_size", 500),
        chunk_overlap=config.get("chunk_overlap", 50)
    )


"""
Document Processor
Handles uploaded files and routes them to the correct pipeline:
  - Structured data (CSV, Excel, JSON) → PostgreSQL + TAG schema
  - Unstructured docs (PDF, TXT, DOCX, MD) → ChromaDB RAG collection

Usage:
    from document_processor import DocumentProcessor
    processor = DocumentProcessor(tag=tag_instance, executor=executor_instance)
    result = processor.process("path/to/file.csv")
    result = processor.process("path/to/report.pdf")
"""

import os
import json
import uuid
import logging
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# File type classification
# ---------------------------------------------------------------------------

STRUCTURED_EXTENSIONS   = {".csv", ".xlsx", ".xls", ".json"}
UNSTRUCTURED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".markdown"}


def classify_file(file_path: str) -> str:
    """
    Classify a file as 'structured' or 'unstructured'.

    Returns:
        'structured'   → CSV, Excel, JSON  (goes to Postgres + TAG)
        'unstructured' → PDF, TXT, DOCX, MD (goes to ChromaDB RAG)
        'unsupported'  → anything else
    """
    ext = Path(file_path).suffix.lower()
    if ext in STRUCTURED_EXTENSIONS:
        return "structured"
    if ext in UNSTRUCTURED_EXTENSIONS:
        return "unstructured"
    return "unsupported"


# ---------------------------------------------------------------------------
# Structured data handlers
# ---------------------------------------------------------------------------

class StructuredFileLoader:
    """Loads structured files (CSV, Excel, JSON) into a pandas DataFrame."""

    def load(self, file_path: str):
        """
        Load any supported structured file into a DataFrame.
        Returns (DataFrame, table_name_suggestion).
        """
        import pandas as pd

        ext  = Path(file_path).suffix.lower()
        stem = Path(file_path).stem.lower().replace(" ", "_").replace("-", "_")

        if ext == ".csv":
            df = pd.read_csv(file_path)
        elif ext in {".xlsx", ".xls"}:
            df = pd.read_excel(file_path)
        elif ext == ".json":
            df = pd.read_json(file_path)
        else:
            raise ValueError(f"Unsupported structured format: {ext}")

        # Sanitize column names
        df.columns = [
            c.lower().strip().replace(" ", "_").replace("-", "_")
            for c in df.columns
        ]

        logger.info(f"Loaded {len(df)} rows, {len(df.columns)} columns from {file_path}")
        return df, stem

    def infer_postgres_type(self, dtype) -> str:
        """Map pandas dtype to PostgreSQL type."""
        import pandas as pd
        dtype_str = str(dtype)
        if "int" in dtype_str:
            return "BIGINT"
        if "float" in dtype_str:
            return "DOUBLE PRECISION"
        if "bool" in dtype_str:
            return "BOOLEAN"
        if "datetime" in dtype_str:
            return "TIMESTAMP"
        if "date" in dtype_str:
            return "DATE"
        return "TEXT"

    def create_table_and_insert(
        self,
        df,
        table_name: str,
        engine,           # SQLAlchemy engine (admin connection, not read-only)
        if_exists: str = "replace"
    ) -> bool:
        """
        Write DataFrame to PostgreSQL and grant read access to ai_readonly.
        Uses pandas to_sql for simplicity.
        """
        from sqlalchemy import text

        try:
            # Write table (pandas handles CREATE TABLE automatically)
            df.to_sql(
                table_name,
                engine,
                if_exists=if_exists,
                index=False,
                method="multi",
                chunksize=500
            )

            # Grant read access to the read-only role
            with engine.connect() as conn:
                conn.execute(text(
                    f"GRANT SELECT ON TABLE {table_name} TO ai_readonly;"
                ))
                conn.commit()

            logger.info(f"Table '{table_name}' created/updated with {len(df)} rows")
            return True

        except Exception as e:
            logger.error(f"Failed to write table '{table_name}': {e}")
            return False

    def build_table_description(self, df, table_name: str, file_path: str):
        """
        Build a TableDescription object for the TAG layer so the
        SQL agent knows this table exists.
        """
        from layers.layer3_tag import TableDescription

        columns = []
        sample_values = {}

        for col in df.columns:
            pg_type = self.infer_postgres_type(df[col].dtype)
            columns.append({
                "name":        col,
                "type":        pg_type,
                "description": f"Column from uploaded file {Path(file_path).name}"
            })
            # Pick a non-null sample value
            non_null = df[col].dropna()
            if len(non_null) > 0:
                sample_values[col] = str(non_null.iloc[0])

        return TableDescription(
            table_name=table_name,
            description=(
                f"Uploaded structured dataset from '{Path(file_path).name}'. "
                f"Contains {len(df)} rows and {len(df.columns)} columns."
            ),
            columns=columns,
            relationships=[],
            sample_values=sample_values
        )


# ---------------------------------------------------------------------------
# Unstructured document handlers
# ---------------------------------------------------------------------------

class UnstructuredFileLoader:
    """Extracts text from PDF, TXT, DOCX, MD files."""

    def load(self, file_path: str) -> str:
        """Extract raw text from the file. Returns text string."""
        ext = Path(file_path).suffix.lower()

        if ext == ".txt" or ext in {".md", ".markdown"}:
            return self._load_text(file_path)
        elif ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported unstructured format: {ext}")

    def _load_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf (install: pip install pypdf)."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF support. Install it with:\n"
                "  pip install pypdf"
            )

    def _load_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx (install: pip install python-docx)."""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. Install it with:\n"
                "  pip install python-docx"
            )

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50       # FIX: renamed overlap → chunk_overlap to match caller
    ) -> List[str]:
        """
        Split text into overlapping chunks for embedding.
        Simple word-boundary chunking.
        """
        words  = text.split()
        chunks = []
        start  = 0

        while start < len(words):
            end   = min(start + chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk)
            if end == len(words):
                break
            start = end - chunk_overlap  # overlap for context continuity

        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks


# ---------------------------------------------------------------------------
# Main DocumentProcessor
# ---------------------------------------------------------------------------

class DocumentProcessor:
    """
    Main entry point for file processing.

    Requires:
        tag      - TAGRetrieval instance (Layer 3)
        executor - SecureExecutionSandbox instance (Layer 5) for test_connection()

    For structured files, also needs:
        admin_engine - SQLAlchemy engine with write permissions (NOT the read-only one)
    """

    def __init__(
        self,
        tag,
        executor=None,
        admin_db_url: str = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        """
        Args:
            tag:           TAGRetrieval instance
            executor:      SecureExecutionSandbox (used only for health check)
            admin_db_url:  SQLAlchemy URL with write access for structured uploads
                           e.g. "postgresql://postgres:password@localhost/postgres"
            chunk_size:    Words per chunk for RAG
            chunk_overlap: Overlapping words between chunks
        """
        self.tag          = tag
        self.executor     = executor
        self.chunk_size   = chunk_size
        self.chunk_overlap = chunk_overlap

        self._structured_loader   = StructuredFileLoader()
        self._unstructured_loader = UnstructuredFileLoader()

        # Admin engine (write access) for structured uploads
        self._admin_engine = None
        if admin_db_url:
            from sqlalchemy import create_engine
            self._admin_engine = create_engine(admin_db_url, pool_pre_ping=True)
            logger.info("Admin DB engine initialized for structured uploads")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Main method. Auto-detects file type and processes accordingly.

        Returns a result dict:
        {
            "success":   bool,
            "file_type": "structured" | "unstructured" | "unsupported",
            "file_name": str,
            "message":   str,
            # structured only:
            "table_name": str,
            "row_count":  int,
            "columns":    List[str],
            # unstructured only:
            "chunk_count": int,
            "doc_ids":     List[str],
        }
        """
        file_path = str(file_path)
        file_name = Path(file_path).name
        file_type = classify_file(file_path)

        logger.info(f"Processing '{file_name}' as {file_type}")

        if file_type == "structured":
            return self._process_structured(file_path, file_name)
        elif file_type == "unstructured":
            return self._process_unstructured(file_path, file_name)
        else:
            return {
                "success":   False,
                "file_type": "unsupported",
                "file_name": file_name,
                "message":   (
                    f"Unsupported file type: {Path(file_path).suffix}. "
                    f"Supported: {STRUCTURED_EXTENSIONS | UNSTRUCTURED_EXTENSIONS}"
                )
            }

    def process_many(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple files. Returns list of result dicts."""
        return [self.process(p) for p in file_paths]

    def list_loaded_schemas(self) -> List[str]:
        """List all table names currently in the TAG schema collection."""
        try:
            count = self.tag.schema_collection.count()
            if count == 0:
                return []
            results = self.tag.schema_collection.get()
            return results.get("ids", [])
        except Exception:
            return []

    def list_loaded_documents(self) -> List[Dict[str, str]]:
        """List all document chunks in the TAG docs collection."""
        try:
            count = self.tag.docs_collection.count()
            if count == 0:
                return []
            results = self.tag.docs_collection.get()
            ids       = results.get("ids", [])
            metadatas = results.get("metadatas", [{}] * len(ids))
            return [
                {"id": id_, "file_name": m.get("file_name", "unknown")}
                for id_, m in zip(ids, metadatas)
            ]
        except Exception:
            return []

    # ------------------------------------------------------------------
    # Structured processing
    # ------------------------------------------------------------------

    def _process_structured(self, file_path: str, file_name: str) -> Dict[str, Any]:
        try:
            # 1. Load into DataFrame
            df, table_name = self._structured_loader.load(file_path)

            # 2. Write to Postgres (if admin engine available)
            db_written = False
            if self._admin_engine:
                db_written = self._structured_loader.create_table_and_insert(
                    df, table_name, self._admin_engine
                )
                if not db_written:
                    logger.warning(
                        f"Could not write '{table_name}' to DB. "
                        "Schema will still be added to TAG so the LLM knows about it."
                    )
            else:
                logger.warning(
                    "No admin_db_url provided — skipping DB write. "
                    "Schema will be added to TAG only."
                )

            # 3. Add schema to TAG so SQL agent can query it
            table_desc = self._structured_loader.build_table_description(
                df, table_name, file_path
            )
            self.tag.add_schema(table_desc)

            return {
                "success":    True,
                "file_type":  "structured",
                "file_name":  file_name,
                "table_name": table_name,
                "row_count":  len(df),
                "columns":    list(df.columns),
                "db_written": db_written,
                "message":    (
                    f"Loaded {len(df)} rows into table '{table_name}' "
                    f"and added schema to TAG."
                    + ("" if db_written else " (DB write skipped — no admin connection)")
                )
            }

        except Exception as e:
            logger.error(f"Structured processing failed for '{file_name}': {e}")
            return {
                "success":   False,
                "file_type": "structured",
                "file_name": file_name,
                "message":   f"Failed: {str(e)}"
            }

    # ------------------------------------------------------------------
    # Unstructured processing
    # ------------------------------------------------------------------

    def _process_unstructured(self, file_path: str, file_name: str) -> Dict[str, Any]:
        try:
            # 1. Extract text
            text = self._unstructured_loader.load(file_path)

            if not text.strip():
                return {
                    "success":   False,
                    "file_type": "unstructured",
                    "file_name": file_name,
                    "message":   "File appears to be empty or could not extract text."
                }

            # 2. Chunk text
            chunks = self._unstructured_loader.chunk_text(
                text,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )

            # 3. Store each chunk in ChromaDB docs collection
            # Use file hash + chunk index as stable ID
            file_hash = hashlib.md5(file_name.encode()).hexdigest()[:8]
            doc_ids   = []

            for i, chunk in enumerate(chunks):
                doc_id = f"{file_hash}_chunk_{i}"
                self.tag.add_document(
                    doc_id=doc_id,
                    content=chunk,
                    metadata={
                        "file_name":  file_name,
                        "file_path":  file_path,
                        "chunk_index": str(i),
                        "total_chunks": str(len(chunks)),
                        "file_type":  Path(file_path).suffix.lower()
                    }
                )
                doc_ids.append(doc_id)

            return {
                "success":     True,
                "file_type":   "unstructured",
                "file_name":   file_name,
                "chunk_count": len(chunks),
                "doc_ids":     doc_ids,
                "char_count":  len(text),
                "message":     (
                    f"Extracted {len(text):,} characters from '{file_name}', "
                    f"split into {len(chunks)} chunks, stored in RAG collection."
                )
            }

        except Exception as e:
            logger.error(f"Unstructured processing failed for '{file_name}': {e}")
            return {
                "success":   False,
                "file_type": "unstructured",
                "file_name": file_name,
                "message":   f"Failed: {str(e)}"
            }


# ---------------------------------------------------------------------------
# Factory function
# ---------------------------------------------------------------------------

def create_document_processor(
    tag,
    executor=None,
    config: Dict[str, Any] = None
) -> DocumentProcessor:
    """
    Create a DocumentProcessor from config or environment variables.

    Config keys (config.yaml):
        db_host, db_port, db_name
        admin_db_user, admin_db_password   <- preferred write-access user
        db_user, db_password               <- fallback
        chunk_size, chunk_overlap

    Environment variable fallbacks:
        DB_HOST, DB_PORT, DB_NAME
        ADMIN_DB_USER, ADMIN_DB_PASSWORD
        DB_USER, DB_PASSWORD
    """
    config = config or {}

    db_host   = config.get("db_host",  os.getenv("DB_HOST",  "localhost"))
    db_port   = config.get("db_port",  int(os.getenv("DB_PORT", "5432")))
    db_name   = config.get("db_name",  os.getenv("DB_NAME",  "postgres"))

    admin_user = (
        config.get("admin_db_user")
        or os.getenv("ADMIN_DB_USER")
        or config.get("db_user")
        or os.getenv("DB_USER", "postgres")
    )
    admin_pass = (
        config.get("admin_db_password")
        or os.getenv("ADMIN_DB_PASSWORD")
        or config.get("db_password")
        or os.getenv("DB_PASSWORD", "")
    )

    admin_db_url = None
    if admin_user and admin_pass:
        admin_db_url = (
            f"postgresql://{admin_user}:{admin_pass}"
            f"@{db_host}:{db_port}/{db_name}"
        )
        logger.info(f"Admin DB URL set for user '{admin_user}' on '{db_name}'")
    else:
        logger.warning(
            "No admin DB credentials found. Structured files will be added to "
            "TAG schema only, not written to PostgreSQL.\n"
            "Fix: add to config.yaml:\n"
            "  admin_db_user: postgres\n"
            "  admin_db_password: yourpassword\n"
            "Or set env vars: ADMIN_DB_USER / ADMIN_DB_PASSWORD"
        )

    return DocumentProcessor(
        tag=tag,
        executor=executor,
        admin_db_url=admin_db_url,
        chunk_size=config.get("chunk_size", 500),
        chunk_overlap=config.get("chunk_overlap", 50)
    )


# ---------------------------------------------------------------------------
# Quick test / demo
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import tempfile, csv

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # We need a TAG instance to test
    from layers.layer3_tag import TAGRetrieval
    tag = TAGRetrieval()

    processor = DocumentProcessor(tag=tag)

    # --- Test 1: CSV ---
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
        writer = csv.writer(f)
        writer.writerow(["product_name", "sales", "region"])
        writer.writerow(["Widget A", 1500, "North America"])
        writer.writerow(["Widget B", 900, "Europe"])
        csv_path = f.name

    result = processor.process(csv_path)
    print("\nCSV result:")
    print(json.dumps({k: v for k, v in result.items() if k != "doc_ids"}, indent=2))

    # --- Test 2: TXT ---
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("This is a sample policy document. " * 100)
        txt_path = f.name

    result = processor.process(txt_path)
    print("\nTXT result:")
    print(json.dumps({k: v for k, v in result.items() if k != "doc_ids"}, indent=2))

    # --- Test 3: Unsupported ---
    result = processor.process("some_file.pptx")
    print("\nUnsupported result:")
    print(json.dumps(result, indent=2))
